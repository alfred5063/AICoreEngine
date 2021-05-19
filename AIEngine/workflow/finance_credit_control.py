# Declare Python libraries needed for this script
import pandas as pd
import xlrd
import numpy as np
import os
import docx
import datetime
from datetime import date
from docx import Document   #Import the Document
from docx.shared import Inches   #Import Inches, used to set width, height etc
from docx.shared import Pt   #Import Pt, used to font etc
from docx.enum.table import WD_TABLE_ALIGNMENT
from num2words import num2words
from decimal import Decimal
from docx2pdf import convert
from directory.movefile import copy_file_to_archived, move_file_to_result_medi
from directory.get_filename_from_path import get_file_name
from utils.Session import session
from utils.audit_trail import *
from utils.logging import logging
from utils.notification import send
from directory.directory_setup import prepare_directories

#excelf = r'C:\Users\Asus\Desktop\Credit Control\Ops disbursement_ MEM_ 2019.xlsx'
#docf = r'C:\Users\Asus\Desktop\Credit Control\LOD-OpsInv-v7-190409(edited).docx'

def process_credit_control(base, excel_file, doc_file):
  #base = session.base(taskid, guid, email, step_name, filename=excel_file)
  head, tail = get_file_name(excel_file)
  prepare_directories(head, base)
  result = populate_data_medi(excel_file, doc_file, base)
  return result

def process_credit_control(excelf, docf, base):
  path, filename = os.path.split(docf)
  audit_log("Credit control processing", "Completed...", base)
  try:

    wb = xlrd.open_workbook(excelf)
    filesheet = wb.sheet_names()[0]
    print(filesheet)
    ###Change nrows to read records!!
    row_data = pd.read_excel(excelf, sheet_name = filesheet, header = 4 ,nrows=10)
    row_data = row_data[pd.notnull(row_data['Client'])] # Only take records where Date is the latest date
    row_data.columns
    
    for index, row in row_data.iterrows():
      
      dc_no= row['Disbursement claims no']
      policy_no= row['Policy']
      hospital= row['Hospital ']
      medical_expenses= row['Insurance']
      medical_claim_statement_sum= row['Total ']
      insurer_name= row['Client']
      insurer_address= row['Address']
      start_date= row['Admission date ']
      end_date= row['Discharge date ']
      bill_no= row['Bill no.']
      short_fall_difference =row['AA Over Guaranteed']
      file_no =row['File no']
      bord_no =row['Bord']
      patient =row['Patient']

      
      if type(start_date) == datetime.datetime:
        start_date = str(start_date.strftime("%x"))
      else:
        start_date = str(start_date)
        
      if type(end_date) == datetime.datetime:
        end_date = str(end_date.strftime("%x"))
      else:
        end_date = str(end_date)
      
      
      audit_log("Credit control word creation", "Completed...", base)
      doc = docx.Document(docf)

      TWOPLACES = Decimal(10) ** -2       # same as Decimal('0.01')
      medical_claim_statement_sum = str(Decimal(medical_claim_statement_sum).quantize(TWOPLACES))
      medical_expenses = str(Decimal(medical_expenses).quantize(TWOPLACES))
      short_fall_difference = str(Decimal(short_fall_difference).quantize(TWOPLACES))
      medical_claim_statement_sum_in_word  = num2words(medical_claim_statement_sum, to='currency')
      medical_claim_statement_sum_in_word = medical_claim_statement_sum_in_word.replace("euro", "")
      medical_expenses_in_word  = num2words(medical_expenses, to='currency')
      medical_expenses_in_word = medical_expenses_in_word.replace("euro", "")
      short_fall_difference_in_word  = num2words(short_fall_difference, to='currency')
      short_fall_difference_in_word = short_fall_difference_in_word.replace("euro", "")

      today = date.today()
      today_date = today.strftime("%d/%m/%Y")
      doc.paragraphs[0].runs[1].text = today_date
      test_run=doc.paragraphs[1].add_run()
      test_run.text = dc_no
      
      doc.paragraphs[3].runs[14].text  = ' to be recorded...'
      doc.paragraphs[13].runs[1].text  = short_fall_difference
      doc.paragraphs[15].runs[1].text = insurer_name
    
      test_run=doc.paragraphs[15].add_run()
      test_run.text = str(insurer_address)
      font = test_run.font
      font.size = Pt(12)
      doc.paragraphs[17].runs[2].text  = hospital
      doc.paragraphs[17].runs[5].text = start_date
      doc.paragraphs[17].runs[8].text = end_date
      doc.paragraphs[17].runs[11].text = ' '+medical_expenses
      doc.paragraphs[17].runs[14].text = medical_expenses_in_word
     
      doc.paragraphs[19].runs[2].text = insurer_name
      doc.paragraphs[19].runs[5].text = policy_no
      doc.paragraphs[19].runs[7].text = hospital
      doc.paragraphs[19].runs[13].text = end_date
      doc.paragraphs[19].runs[10].text  = str(bill_no)
      doc.paragraphs[19].runs[17].text = doc.paragraphs[19].runs[17].text +medical_claim_statement_sum
      doc.paragraphs[19].runs[21].text = short_fall_difference
      doc.paragraphs[19].runs[23].text = short_fall_difference_in_word

      doc.paragraphs[21].runs[1].text = start_date

      doc.paragraphs[25].runs[2].text = short_fall_difference

      doc.paragraphs[50].runs[1].text = patient

      #TABLE
      doc.tables[0].cell(0,2).text = insurer_name
      doc.tables[0].cell(1,2).text = policy_no
      doc.tables[0].cell(2,2).text = hospital
      doc.tables[0].cell(3,2).text = doc.tables[0].cell(3,2).text + medical_expenses
      doc.tables[0].cell(4,2).text = doc.tables[0].cell(4,2).text + medical_claim_statement_sum


      doc.tables[1].cell(1,0).text = file_no
      doc.tables[1].cell(1,0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
      doc.tables[1].cell(1,1).text = dc_no
      doc.tables[1].cell(1,1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
      doc.tables[1].cell(1,2).text = bord_no
      doc.tables[1].cell(1,2).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
      doc.tables[1].cell(1,3).text = today_date
      doc.tables[1].cell(1,3).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER


      
      doc.tables[2].cell(2,3).text = doc.tables[2].cell(2,3).text +short_fall_difference
      doc.tables[2].cell(2,2).text = 'Being charges incurred at ' + hospital
      doc.tables[2].cell(2,5).text = 'on '+ start_date+ ' to ' +end_date
      doc.tables[2].cell(6,2).text = 'I/c No: ' + 'to be recorded...'
      doc.tables[2].cell(5,2).text=  'Patient\'s name: ' + patient
      doc.tables[2].cell(15,2).text = 'Ringgit Malaysia '+short_fall_difference_in_word+' only)'
      doc.tables[2].cell(17,3).text = doc.tables[2].cell(17,3).text + short_fall_difference
      
      doc.tables[2].cell(17,3).paragraphs[0].runs[0].font.bold = True
      
      
      doc.save(path+'\\'+dc_no+'.docx')
      convert(path+'\\'+dc_no+'.docx', path+'\\'+dc_no+'.pdf')

      audit_log("Credit control process completed", "Completed...", base)
      return 'Success'
  except Exception as error:
    print("Exception error occur: {0}".format(error))
    logging('populate_data_medi', error, base)
    return 'Fail'
    

