#!/usr/bin/python
# FINAL SCRIPT updated as of 06st July 2020
# Workflow - Finance Credit Control

# Declare Python libraries needed for this script
import shutil
import sys
import os
import glob
import json as json
import requests
import pandas as pd
import xlrd
import numpy as np
import os
import docx
import datetime
import automagica
from datetime import datetime
from pandas.io.json import json_normalize
from utils.Session import session
from utils.guid import get_guid
from utils.notification import send
from utils.audit_trail import audit_log
from utils.logging import logging
from connector.connector import *
from directory.directory_setup import prepare_directories
from datetime import date
from docx import Document   #Import the Document
from docx.shared import Inches   #Import Inches, used to set width, height etc
from docx.shared import Pt   #Import Pt, used to font etc
from docx.enum.table import WD_TABLE_ALIGNMENT
from num2words import num2words
from decimal import Decimal
from docx2pdf import convert
from automagic.automagica_creditcontrol import *
from transformation.adm_fill_bdx import Check_post_or_type
from directory.movefile import *

def process_cc(excelf, docf, creditcontrol_base):

  path, filename = os.path.split(docf)

  try:

    wb = xlrd.open_workbook(excelf)
    filesheet = wb.sheet_names()[0]
    row_data = pd.read_excel(excelf, sheet_name = filesheet, header = 4)
    row_data = row_data[pd.notnull(row_data['Client'])]
    
    for index, row in row_data.iterrows():

      try:
        icno = row['IC Number']
        dc_no = row['Disbursement claims no']
        policy_no = row['Policy']
        hospital = row['Hospital ']
        medical_expenses = row['Amount         (RM)']
        medical_claim_statement_sum = row['Cheq amounts']
        insurer_name = row['Client']
        patient_address = row['Patient Address']
        insurer_address = row['Insurer Address']

        if isinstance(row['Admission date '], str) == True:
          start_date = row['Admission date ']
        else:
          start_date = row['Admission date '].strftime("%d/%m/%Y")

        if isinstance(row['Discharge date '], str) == True:
          end_date = row['Discharge date ']
        else:
          end_date = row['Discharge date '].strftime("%d/%m/%Y")

        if isinstance(row['Letter Date'], str) == True:
          invoice_date = row['Letter Date']
        else:
          invoice_date = row['Letter Date'].strftime("%d/%m/%Y")

        bill_no = row['Hospital Bill']
        short_fall_difference = row['AA Over Guaranteed']
        file_no = row['File no']
        bord_no = row['Bord']
        patient = row['Patient']
        reasons = row['Reasons']

        # Calling MARC to get IC and Address
        if "-" in str(row['File no']):
          caseid = str(row['File no']).split("-")
          caseid = caseid[1]
        else:
          caseid = row['File no']

        if str(insurer_name) != 'HLA':
          type = str(Check_post_or_type(caseid))
          browser = login_Marc(creditcontrol_base)
          if type == 'POST':
            if icno == '':
              navigate_page(browser, 'inpatient', 'csucases', 'inpatient_cases_csu_search')
              icno = search_post_member(browser, caseid, creditcontrol_base)
            else:
              pass
            if patient_address == '':
              try:
                patient_address = get_post_address(browser)
                print("1")
              except Exception:
                patient_address = " "
                print("2")
            else:
              pass
          elif type == 'CASHLESS':
            if icno == '':
              navigate_page(browser, 'inpatient', 'cases', 'inpatient_cases_search')
              icno = search_adm_member(browser, caseid, creditcontrol_base)
            else:
              pass
            if patient_address == '':
              try:
                patient_address = get_adm_address(browser)
                print("3")
              except Exception:
                patient_address = " "
                print("4")
            else:
              pass
          else:
            icno = "Unable to extract member's IC number from MARC. Please add manually."
            patient_address = "Unable to extract member's mailing address from MARC. Please add manually."
            print("5")
          browser.close()
        else:
          print("6")
          icno = ''
          patient_address = ''

        doc = docx.Document(docf)

        TWOPLACES = Decimal(10) ** -2
        medical_claim_statement_sum = str(Decimal(medical_claim_statement_sum).quantize(TWOPLACES))
        medical_expenses = str(Decimal(medical_expenses).quantize(TWOPLACES))
        short_fall_difference = str(Decimal(short_fall_difference).quantize(TWOPLACES))
        medical_claim_statement_sum_in_word = num2words(medical_claim_statement_sum, to = 'currency')
        medical_claim_statement_sum_in_word = medical_claim_statement_sum_in_word.replace("euro", "")
        medical_expenses_in_word = num2words(medical_expenses, to = 'currency')
        medical_expenses_in_word = medical_expenses_in_word.replace("euro", "")
        short_fall_difference_in_word = num2words(short_fall_difference, to = 'currency')
        short_fall_difference_in_word = short_fall_difference_in_word.replace("euro", "")

        today = date.today()
        today_date = today.strftime("%d/%m/%Y")
        doc.paragraphs[0].runs[1].text = today_date
        test_run = doc.paragraphs[1].add_run()
        test_run.text = dc_no

        # First Page
        doc.paragraphs[3].runs[14].text = str(icno)
        doc.paragraphs[4].runs[8].text = str(patient_address)
        doc.paragraphs[13].runs[1].text = short_fall_difference
        doc.paragraphs[15].runs[1].text = insurer_name
        test_run = doc.paragraphs[15].add_run()
        test_run.text = str(insurer_address)
        font = test_run.font
        font.size = Pt(12)
        doc.paragraphs[17].runs[2].text = hospital
        doc.paragraphs[17].runs[5].text = str(start_date)
        doc.paragraphs[17].runs[8].text = str(end_date)
        doc.paragraphs[17].runs[12].text = ' ' + str(medical_expenses)
        doc.paragraphs[17].runs[15].text = medical_expenses_in_word
        doc.paragraphs[19].runs[2].text = insurer_name
        doc.paragraphs[19].runs[5].text = policy_no
        doc.paragraphs[19].runs[7].text = hospital
        doc.paragraphs[19].runs[10].text = ' ' + str(bill_no)
        doc.paragraphs[19].runs[12].text = str(end_date)
        doc.paragraphs[19].runs[15].text = str(medical_claim_statement_sum) + ' '
        doc.paragraphs[19].runs[19].text = str(short_fall_difference)
        doc.paragraphs[19].runs[22].text = str(short_fall_difference_in_word.replace(' ,', ',')) + ' '
        doc.paragraphs[21].runs[1].text = str(start_date)
        doc.paragraphs[25].runs[2].text = short_fall_difference

        # Second Page
        doc.paragraphs[44].runs[0].text = str(patient_address)
        doc.paragraphs[50].runs[1].text = patient
        doc.tables[0].cell(0,2).text = insurer_name
        doc.tables[0].cell(1,2).text = policy_no
        doc.tables[0].cell(2,2).text = hospital
        doc.tables[0].cell(3,2).text = str(doc.tables[0].cell(3,2).text) + str(medical_expenses)
        doc.tables[0].cell(4,2).text = str(doc.tables[0].cell(4,2).text) + str(medical_claim_statement_sum)
        doc.tables[1].cell(1,0).text = str(file_no)
        doc.tables[1].cell(1,0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
        doc.tables[1].cell(1,1).text = dc_no
        doc.tables[1].cell(1,1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
        doc.tables[1].cell(1,2).text = bord_no
        doc.tables[1].cell(1,2).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
        doc.tables[1].cell(1,3).text = str(invoice_date)
        doc.tables[1].cell(1,3).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
        doc.tables[2].cell(2,3).text = str(doc.tables[2].cell(2,3).text) + str(short_fall_difference)
        doc.tables[2].cell(2,2).text = 'Being charges incurred at ' + str(hospital)
        doc.tables[2].cell(2,5).text = 'on '+ str(start_date) + ' to ' + str(end_date)
        doc.tables[2].cell(5,2).text = 'Patient\'s name: ' + str(patient)
        doc.tables[2].cell(6,2).text = 'I/c No: ' + str(icno)
        doc.tables[2].cell(6,4).text = 'Reasons: ' + str(reasons)
        doc.tables[2].cell(15,2).text = 'Ringgit Malaysia ' + str(short_fall_difference_in_word) + ' only)'
        doc.tables[2].cell(17,3).text = str(doc.tables[2].cell(17,3).text) + str(short_fall_difference)
        doc.tables[2].cell(17,3).paragraphs[0].runs[0].font.bold = True

        try:
          file_word = str(path) + '\\' + str(dc_no) + '.docx'
          doc.save(file_word)

          file_pdf = str(path) + '\\' + str(dc_no) + '.pdf'
          #convert(file_word, file_pdf)

          print("Credit Control - Generate Member Letter - Successfully generate word document and pdf for [ %s ]. Processing another record. Please wait." %file_no)
          audit_log("Credit Control - Generate Member Letter - Successfully generate word document and pdf for [ %s ]. Processing another record. Please wait." %file_no, "Completed...", creditcontrol_base)
        except Exception as error:
          audit_log("Credit Control - Generate Member Letter - Unsuccessful generating word document and pdf for [ %s ]. Please check manually. Proceed to another record. Please wait." %file_no, "Completed...", creditcontrol_base)
          pass

        time.sleep(5)

        move_file_to_result_ar(file_word, creditcontrol_base)
        #move_file_to_result_ar(file_pdf, creditcontrol_base)

      except Exception as error:
        print("Credit Control - Generate Member Letter - Unsuccessful generating word document and pdf for [ %s ]. Please check manually." %file_no)
        audit_log("Credit Control - Generate Member Letter - Unsuccessful generating word document and pdf for [ %s ]. Please check manually. Proceed to another record. Please wait." %file_no, "Completed...", creditcontrol_base)
        pass
      continue

  except Exception as error:
    logging('Credit Control - Generate Member Letter.', error, creditcontrol_base)
